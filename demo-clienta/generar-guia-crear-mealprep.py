from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "guia-crear-mealprep.docx"

INK = "21160E"
BEET = "8E263C"
BEET_SOFT = "F5E7EA"
PAPER = "FFF9F1"
MUTED = "756A60"
LINE = "D9C9C3"


def set_run_font(run, name="Arial", size=10, color=INK, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def shade(cell, color):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE):
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for side in ("top", "left", "bottom", "right"):
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "6")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def prevent_row_split(row):
    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    properties.append(cant_split)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width
            tc_properties = row.cells[index]._tc.get_or_add_tcPr()
            cell_width = tc_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                tc_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(int(width.inches * 1440)))
            cell_width.set(qn("w:type"), "dxa")


def style_paragraph(paragraph, before=0, after=4, line=1.18):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_text(paragraph, text, size=10, color=INK, bold=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return run


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, before=14 if level == 1 else 10, after=6, line=1)
    add_text(paragraph, text, size=14 if level == 1 else 11, color=BEET, bold=True)
    return paragraph


def add_field_table(document, rows):
    table = document.add_table(rows=1, cols=2)
    set_table_widths(table, [Inches(1.875), Inches(4.625)])
    headers = table.rows[0].cells
    for cell, text in zip(headers, ("Campo", "Qué significa y cómo completarlo")):
        shade(cell, BEET)
        set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
        set_cell_border(cell, color=BEET)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        style_paragraph(paragraph, before=0, after=0, line=1.1)
        add_text(paragraph, text, size=9.5, color="FFFFFF", bold=True)

    for label, description in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for cell in cells:
            shade(cell, PAPER)
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        label_paragraph = cells[0].paragraphs[0]
        style_paragraph(label_paragraph, before=0, after=0, line=1.12)
        add_text(label_paragraph, label, size=9.4, color=INK, bold=True)
        detail_paragraph = cells[1].paragraphs[0]
        style_paragraph(detail_paragraph, before=0, after=0, line=1.12)
        add_text(detail_paragraph, description, size=9.25, color=INK)

    return table


def add_callout(document, title, body):
    table = document.add_table(rows=1, cols=1)
    set_table_widths(table, [Inches(6.5)])
    cell = table.cell(0, 0)
    shade(cell, BEET_SOFT)
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    set_cell_border(cell, color="D9A9B4")
    title_paragraph = cell.paragraphs[0]
    style_paragraph(title_paragraph, before=0, after=4, line=1)
    add_text(title_paragraph, title, size=10, color=BEET, bold=True)
    body_paragraph = cell.add_paragraph()
    style_paragraph(body_paragraph, before=0, after=0, line=1.14)
    add_text(body_paragraph, body, size=10, color=INK)


def add_steps_table(document):
    table = document.add_table(rows=1, cols=2)
    set_table_widths(table, [Inches(0.65), Inches(5.85)])
    headers = table.rows[0].cells
    for cell, text in zip(headers, ("Paso", "Revisión antes de guardar")):
        shade(cell, BEET)
        set_cell_margins(cell, top=100, start=130, bottom=100, end=130)
        set_cell_border(cell, color=BEET)
        paragraph = cell.paragraphs[0]
        style_paragraph(paragraph, before=0, after=0, line=1.1)
        add_text(paragraph, text, size=9.5, color="FFFFFF", bold=True)
    steps = [
        ("1", "Completar nombre, tipo, precio y orden."),
        ("2", "Escribir descripción, beneficios, ingredientes y alérgenos."),
        ("3", "Cargar foto principal y, si corresponde, foto hover."),
        ("4", "Revisar que Activo esté encendido solo cuando el producto esté listo."),
        ("5", "Guardar y revisar el producto en la tienda pública antes de comunicarlo.")
    ]
    for number, action in steps:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for cell in cells:
            shade(cell, PAPER)
            set_cell_margins(cell, top=80, start=130, bottom=80, end=130)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        first = cells[0].paragraphs[0]
        first.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(first, before=0, after=0, line=1)
        add_text(first, number, size=9.5, color=BEET, bold=True)
        second = cells[1].paragraphs[0]
        style_paragraph(second, before=0, after=0, line=1.12)
        add_text(second, action, size=9.25)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.92)
    section.right_margin = Inches(0.92)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(paragraph, before=0, after=0, line=1)
    add_text(paragraph, "FULLNESS LAB  |  Guía de creación de Meal Prep", size=8, color=MUTED, bold=True)


def add_title_block(document):
    kicker = document.add_paragraph()
    style_paragraph(kicker, before=0, after=6, line=1)
    add_text(kicker, "FULLNESS LAB  |  BACKOFFICE", size=9, color=BEET, bold=True)

    title = document.add_paragraph()
    style_paragraph(title, before=0, after=4, line=0.98)
    add_text(title, "Guía para crear un Meal Prep", size=23, color=INK, bold=True)

    subtitle = document.add_paragraph()
    style_paragraph(subtitle, before=0, after=13, line=1.16)
    add_text(subtitle, "Referencia rápida para completar cada campo del catálogo y publicar productos con claridad.", size=10.5, color=MUTED)


def build_document():
    document = Document()
    configure_document(document)
    add_title_block(document)
    add_callout(
        document,
        "ORDEN: ¿PARA QUÉ SIRVE?",
        "Define la posición del producto en la tienda. Un número menor aparece primero. Usa saltos de 10 (10, 20, 30) para poder agregar después un producto entre medio, por ejemplo con 15. No modifica precio, stock ni disponibilidad."
    )

    add_heading(document, "Campos de publicación", level=1)
    add_field_table(document, [
        ("Tipo", "Elige Plan para una propuesta completa o Familiar para un plato individual de mayor formato."),
        ("Frecuencia", "Indica si un Plan se vende de forma semanal o mensual. No aplica para productos Familiares."),
        ("Nombre", "Nombre visible en la tienda. Debe explicar rápidamente qué está comprando la persona."),
        ("Slug", "Dirección interna del producto. Usa minúsculas y guiones, sin espacios: plan-semanal-antinflamatorio."),
        ("SKU", "Código interno opcional para identificar el producto en operaciones y reportes."),
        ("Etiqueta", "Frase corta que acompaña el nombre, por ejemplo: 5 meal preps / 1 semana."),
        ("Precio CLP", "Precio final en pesos chilenos. Escribe solo números, sin puntos ni símbolo $."),
        ("Orden", "Prioridad de aparición en la tienda. Los números menores se muestran primero."),
        ("Porciones / duración", "Aclara cuántas porciones incluye o cuánto dura el plan."),
        ("Botón", "Texto de acción para comprar, por ejemplo: Agregar plan semanal."),
        ("Activo", "Publica o esconde el producto sin eliminarlo. Déjalo inactivo mientras aún lo estás preparando.")
    ])

    document.add_page_break()
    add_heading(document, "Contenido, imágenes y nutrición", level=1)
    add_field_table(document, [
        ("Descripción", "Resumen claro y apetitoso del producto. Explica qué incluye y por qué es una buena elección."),
        ("Tags de beneficios", "Beneficios breves, uno por línea: Alto en proteína, Antiinflamatorio, Fuente de fibra."),
        ("Foto principal", "Imagen que se muestra primero en la tienda y la ficha del producto."),
        ("Foto hover", "Imagen alternativa que aparece al pasar el cursor. Úsala para mostrar un detalle o segunda vista."),
        ("Ingredientes", "Ingredientes principales, uno por línea. Escribe de forma simple y fácil de reconocer."),
        ("Alérgenos", "Ingredientes que pueden provocar alergias, uno por línea: pescado, frutos secos, lácteos, etc."),
        ("Descripción nutricional", "Explicación general del aporte del plato para la alimentación de la semana."),
        ("Características nutricionales", "Beneficios concretos, uno por línea: omega 3, fibra vegetal, energía estable."),
        ("Receta resumida", "Breve descripción de la preparación y de la combinación de sabores."),
        ("Detalle nutricional", "Texto más completo para explicar por qué el producto funciona nutricionalmente."),
        ("Pasos de receta / preparación", "Instrucciones para calentar o servir. Escribe un paso por línea."),
        ("Datos nutricionales JSON", "Datos estructurados para la ficha. Si no se usarán, dejar {}. No escribir texto normal en este campo.")
    ])

    document.add_page_break()
    add_heading(document, "Platos incluidos en un Plan", level=1)
    add_callout(
        document,
        "PLATOS DEL PLAN",
        "Agrega cada meal prep incluido. Cada uno puede llevar nombre, etiqueta, descripción, fotos, beneficios, ingredientes, información nutricional y alérgenos. Esto permite explicar el contenido real del plan antes de comprar."
    )

    add_heading(document, "Antes de guardar", level=1)
    add_steps_table(document)

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
