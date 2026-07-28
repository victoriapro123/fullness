from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SCREENSHOTS = ROOT / "screenshots"
OUTPUT = ROOT / "entregables"
OUTPUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUTPUT / "informe-qa-backoffice-cecilia-2026-07-28.docx"

# Design selection required by the document skill:
# preset: compact_reference_guide
# first-page pattern: customer_pack
# named override: restrained Fullness Lab accent palette.
ACCENT = "963650"
ACCENT_LIGHT = "F5EAED"
INK = "241A1B"
MUTED = "746A6B"
GREEN = "416452"
GREEN_LIGHT = "EAF1ED"
AMBER = "8A5A15"
AMBER_LIGHT = "FBF1DF"
RED = "9B1C1C"
RED_LIGHT = "FBEAEA"
BLUE = "2E74B5"
BORDER = "D8CFCA"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGINS):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in margins.items():
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name in ("top", "left", "bottom", "right"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_paragraph(p, before=0, after=6, line=1.25, alignment=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if alignment is not None:
        p.alignment = alignment
    return p


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for sec in doc.sections:
        header = sec.header.paragraphs[0]
        set_paragraph(header, after=0)
        run = header.add_run("FULLNESS LAB  |  INFORME DE QA")
        set_run_font(run, size=8.5, color=ACCENT, bold=True)
        footer = sec.footer.paragraphs[0]
        set_paragraph(footer, after=0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        run = footer.add_run("Uso interno  |  28 de julio de 2026")
        set_run_font(run, size=8.5, color=MUTED)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, before=8, after=5)
    run = p.add_run(text.upper())
    set_run_font(run, size=9, color=ACCENT, bold=True)
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=5)
    run = p.add_run(text)
    set_run_font(run, size=25, color=INK, bold=False)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=14, line=1.2)
    run = p.add_run(text)
    set_run_font(run, size=12, color=MUTED)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, before=0, after=6, color=INK, italic=False):
    p = doc.add_paragraph()
    set_paragraph(p, before=before, after=after)
    run = p.add_run(text)
    set_run_font(run, size=11, color=color, italic=italic)
    return p


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [4050, 2300, 3010])
    headers = ("Prueba", "Resultado", "Evidencia")
    for index, label in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, ACCENT_LIGHT)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        set_paragraph(p, after=0)
        run = p.add_run(label)
        set_run_font(run, size=9.3, color=ACCENT, bold=True)
    for title, status, detail, fill, status_color in rows:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_border(cell)
        set_cell_shading(cells[1], fill)
        values = ((title, INK, False), (status, status_color, True), (detail, MUTED, False))
        for index, (value, color, bold) in enumerate(values):
            p = cells[index].paragraphs[0]
            set_paragraph(p, after=0, line=1.15)
            run = p.add_run(value)
            set_run_font(run, size=9.4, color=color, bold=bold)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, label, text, fill, color):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=color, size="8")
    p = cell.paragraphs[0]
    set_paragraph(p, after=2, line=1.18)
    first = p.add_run(f"{label}  ")
    set_run_font(first, size=9.3, color=color, bold=True)
    second = p.add_run(text)
    set_run_font(second, size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_evidence_image(doc, path, caption):
    p = doc.add_paragraph()
    set_paragraph(p, after=5, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(path), width=Inches(6.35))
    cp = doc.add_paragraph()
    set_paragraph(cp, after=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = cp.add_run(caption)
    set_run_font(run, size=9.2, color=MUTED, italic=True)


def add_label_detail(doc, label, detail):
    p = doc.add_paragraph()
    set_paragraph(p, after=4, line=1.2)
    first = p.add_run(f"{label}: ")
    set_run_font(first, size=10.5, color=INK, bold=True)
    second = p.add_run(detail)
    set_run_font(second, size=10.5, color=INK)


def page_break(doc):
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_document(doc)

    add_kicker(doc, "Validacion operativa")
    add_title(doc, "Informe de QA del backoffice")
    add_subtitle(doc, "Evidencia de acceso administrador, creacion temporal de catalogo y limpieza posterior.")
    add_label_detail(doc, "Cuenta probada", "Perfil administrador de Cecilia; acceso realizado mediante el flujo de autenticacion publicado.")
    add_label_detail(doc, "Entorno", "Backoffice publicado de Fullness Lab. El enlace de autenticacion redirige a fullness-taupe.vercel.app.")
    add_label_detail(doc, "Fecha", "28 de julio de 2026, America/Santiago.")
    add_heading(doc, "Resultado ejecutivo", level=1)
    add_status_table(doc, [
        ("Ingreso al panel y modulos administradores", "APROBADO", "Panel completo visible para Cecilia.", GREEN_LIGHT, GREEN),
        ("Crear mealprep reutilizable con alergenos", "APROBADO", "Registro temporal creado desde la UI.", GREEN_LIGHT, GREEN),
        ("Crear plan semanal y mensual con plato asociado", "APROBADO", "Dos planes temporales guardados desde la UI.", GREEN_LIGHT, GREEN),
        ("Limpiar datos temporales", "APROBADO", "2 planes y 1 mealprep eliminados; consulta final en 0.", GREEN_LIGHT, GREEN),
    ])
    add_callout(
        doc,
        "ALCANCE",
        "La prueba no envio correos, no ejecuto pagos, no creo clientes y no subio archivos a R2, para evitar efectos o costos innecesarios durante QA.",
        AMBER_LIGHT,
        AMBER,
    )

    page_break(doc)
    add_kicker(doc, "Evidencia 01")
    add_title(doc, "Acceso y panel administrador")
    add_subtitle(doc, "El panel se abrio con el perfil Cecilia y mostro los modulos operativos disponibles.")
    add_evidence_image(
        doc,
        SCREENSHOTS / "01-acceso-cecilia-publicado.png",
        "Captura del backoffice publicado. Se observan catalogo, parametros, clientes, contenido y operaciones.",
    )
    add_callout(
        doc,
        "LECTURA",
        "El panel cargado demuestra que Cecilia tiene acceso administrador a la operacion publicada. No se expone ni se modifica ninguna credencial en este informe.",
        GREEN_LIGHT,
        GREEN,
    )

    page_break(doc)
    add_kicker(doc, "Evidencia 02")
    add_title(doc, "Alta de mealprep reutilizable")
    add_subtitle(doc, "La ficha temporal permitio ingresar nombre, descripcion, imagenes, ingredientes, informacion nutricional y alergenos como campo libre.")
    add_evidence_image(
        doc,
        SCREENSHOTS / "02-editor-mealprep.png",
        "Formulario de creacion de plato reutilizable en el backoffice publicado.",
    )
    add_heading(doc, "Flujo ejercitado", level=2)
    add_label_detail(doc, "Registro", "QA 2026-07-28 Bowl temporal.")
    add_label_detail(doc, "Alergenos", "Sesamo y elaborado en cocina compartida.")
    add_label_detail(doc, "Relacion", "El mealprep fue asociado a un plan semanal y a un plan mensual temporales.")

    page_break(doc)
    add_kicker(doc, "Cierre de prueba")
    add_title(doc, "Limpieza y observaciones")
    add_subtitle(doc, "Los registros creados para QA no permanecen en el catalogo publicado.")
    add_heading(doc, "Limpieza verificada", level=1)
    add_status_table(doc, [
        ("Planes QA antes de eliminar", "2 TEMPORALES", "Semanal y mensual.", AMBER_LIGHT, AMBER),
        ("Mealpreps QA antes de eliminar", "1 TEMPORAL", "Bowl reutilizable asociado.", AMBER_LIGHT, AMBER),
        ("Planes QA despues de eliminar", "0 RESTANTES", "Consulta directa a Supabase.", GREEN_LIGHT, GREEN),
        ("Mealpreps QA despues de eliminar", "0 RESTANTES", "Consulta directa a Supabase.", GREEN_LIGHT, GREEN),
    ])
    add_heading(doc, "Observacion de version publicada", level=1)
    add_body(
        doc,
        "La version publicada que se uso para QA conserva etiquetas historicas en el catalogo: \"Meal preps\", \"Platos reutilizables\" y \"Platos familiares\". La taxonomia acordada (Mealpreps, Planes y Mealprep familiar) esta trabajada localmente, pero esta prueba no la despliega ni reemplaza contenido publicado.",
    )
    add_callout(
        doc,
        "CONCLUSION",
        "El acceso administrador de Cecilia y el flujo central de catalogo quedaron validados. Los datos de prueba fueron retirados completamente. Antes de una salida final, corresponde publicar la version que incorpora la nomenclatura acordada y volver a pasar una revision visual corta sobre esa version.",
        ACCENT_LIGHT,
        ACCENT,
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
    print(DOCX_PATH)
