from pathlib import Path

from PIL import Image, ImageOps
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
R2 = ROOT / "r2"
OUTPUT = ROOT / "entregables"
ASSETS = ROOT / "document-assets"
OUTPUT.mkdir(parents=True, exist_ok=True)
ASSETS.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUTPUT / "validacion-recuperacion-mealpreps-cecilia.docx"

# Required design selection:
# - Preset: compact_reference_guide
# - Header pattern: customer_pack
# Named visual override: Fullness brand accent and neutral palette.
ACCENT = "963650"
ACCENT_LIGHT = "F4E8EB"
INK = "241A1B"
MUTED = "746A6B"
CREAM = "FBF8F4"
GREEN = "416452"
GREEN_LIGHT = "EAF1ED"
AMBER = "9A641E"
AMBER_LIGHT = "FBF1DF"
RED = "9B1C1C"
RED_LIGHT = "FBEAEA"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
BLUE_FILL = "E8EEF5"
BORDER = "D8CFCA"
WHITE = "FFFFFF"
PENDING_PAGE_BREAK = False

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


APPLE = {
    "name": "Apple Golden Chicken con puré de zanahoria asada y laurel",
    "status": "COMPLETO Y PRESERVADO",
    "description": "Pollo especiado, puré rústico de camote y hojas verdes.",
    "nutrition_description": (
        "Una preparación reconfortante donde el dulzor natural de la manzana verde "
        "y la zanahoria asada se encuentra con las notas cálidas del curry y la "
        "cúrcuma. Un mealprep equilibrado, lleno de sabor y elaborado con "
        "ingredientes reales para nutrirte desde la raíz."
    ),
    "ingredients": [
        "Pollo",
        "zanahoria",
        "cebolla",
        "manzana verde",
        "leche de almendras",
        "curry",
        "cúrcuma",
        "aceite de oliva extra virgen",
        "sal",
        "pimienta y especias naturales",
    ],
    "allergens": ["Almendras"],
    "benefits": ["Sin gluten", "Ingredientes reales", "Grasas saludables"],
    "highlights": ["Alto en proteína", "Carbohidrato complejo", "Saciedad prolongada"],
    "nutrition": [
        ("Energía", "106 kcal", "265 kcal"),
        ("Proteínas", "16,1 g", "40,2 g"),
        ("Grasas totales", "2,3 g", "5,9 g"),
        ("Carbohidratos disponibles", "3 g", "7,5 g"),
        ("Fibra dietaria", "0,52 g", "1,3 g"),
        ("Sodio", "220 mg", "550 mg"),
    ],
    "portion": "250 g · 1 porción por envase",
    "photos": (
        R2 / "ba8ad241-9696-450f-898f-6e2af142f895.png",
        R2 / "003fc282-a330-46d1-a945-82bf9dd86049.png",
    ),
}

MERLUZA = {
    "name": "Merluza austral sous vide y Crispy Quinoa Salad",
    "status": "COMPLETO Y PRESERVADO",
    "description": (
        "Merluza austral cocinada sous vide y sellada por el lado de la piel, "
        "acompañada de una ensalada de habas, arvejas, espárragos y quinoa con un "
        "cremoso dressing de tahini, limón y jengibre. Una preparación fresca, "
        "rica en proteínas y grasas saludables."
    ),
    "nutrition_description": (
        "Preparación rica en proteínas de alta calidad y fuente natural de omega-3. "
        "Aporta fibra, grasas saludables, vitaminas y minerales provenientes de "
        "legumbres, vegetales y quinoa."
    ),
    "ingredients": [
        "Merluza austral",
        "habas",
        "arvejas",
        "espárragos",
        "pepino",
        "palta",
        "quinoa",
        "cilantro",
        "ciboulette",
        "yogur griego natural",
        "tahini",
        "limón",
        "aceite de oliva extra virgen",
        "miel de palma",
        "aceite de sésamo tostado",
        "jengibre",
        "ralladura de limón",
        "sal marina",
        "pimienta negra",
    ],
    "allergens": [
        "Pescado (merluza austral)",
        "Leche y derivados (yogur griego)",
        "Sésamo (tahini y aceite de sésamo)",
        "Puede contener trazas de frutos secos por elaboración compartida",
    ],
    "benefits": [
        "Alto en proteínas",
        "Alto en fibra",
        "Rico en omega-3",
        "Grasas saludables",
        "Rico en vegetales",
        "Ingredientes reales",
    ],
    "highlights": [
        "Alto en fibra",
        "Rico en omega-3",
        "Alto en proteínas",
        "Fuente de grasas saludables",
        "Rico en vitaminas y minerales",
    ],
    "nutrition": [
        ("Energía", "122 kcal", "623 kcal"),
        ("Proteínas", "9,8 g", "55,2 g"),
        ("Grasas totales", "4,8 g", "24,3 g"),
        ("Carbohidratos disponibles", "9 g", "46,1 g"),
        ("Fibra dietaria", "2,3 g", "11,8 g"),
        ("Sodio", "325 mg", "1.655 mg"),
    ],
    "portion": "511 g · 1 porción por envase",
    "photos": (
        R2 / "0c4c7e19-27fe-452e-b865-f03020469f86.png",
        R2 / "473cdf7c-961a-43d8-9116-3a41167f74b1.png",
    ),
}

CREPES = {
    "name": (
        "Crepes de espinaca y harina de arveja rellenos de salmón ahumado, "
        "queso feta y cebolla acaramelada con miel de dátil"
    ),
    "status": "PARCIAL · REQUIERE VALIDACIÓN",
    "photos": (
        R2 / "810969c1-a948-4a0d-8d31-349d14183f33.png",
        R2 / "5891a268-50ab-46b9-a2cd-a1e2bef8fe01.png",
    ),
}

PLAN_MEALPREPS = [
    "Apple Golden Chicken con puré de zanahoria asada y laurel",
    "Merluza austral sous vide con quinoa crispy salad",
    "Crepes de espinaca y harina de arveja rellenos de salmón ahumado, queso feta y cebolla acaramelada con miel de dátil",
    "Curry de lentejas con salmón sous vide",
    "Hamburguesas de pavo Thai con repollo morado asado y manzana verde",
    "Hamburguesas de garbanzo con salsa pomodoro",
]


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=None, color=INK, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGINS):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[side]))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_border(cell, color=BORDER, size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_borders(table, color=BORDER, size=6):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, color, size)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, BLUE_DARK, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    left = p.add_run("FULLNESS LAB")
    set_run_font(left, size=9, color=ACCENT, bold=True)
    right = p.add_run("   ·   Validación de recuperación")
    set_run_font(right, size=9, color=MUTED)

    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), BORDER)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_kicker(doc, text):
    global PENDING_PAGE_BREAK
    p = doc.add_paragraph()
    if PENDING_PAGE_BREAK:
        p.paragraph_format.page_break_before = True
        PENDING_PAGE_BREAK = False
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    set_run_font(run, size=9, color=ACCENT, bold=True)
    run.font.all_caps = True
    return p


def add_title(doc, text, size=29):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=size, color=INK, bold=True)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    set_run_font(run, size=13, color=MUTED)
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_labeled_paragraph(doc, label, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=11, color=INK, bold=True)
    value_run = p.add_run(text)
    set_run_font(value_run, size=11, color=INK)
    return p


def add_status_callout(doc, label, text, fill, color):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    set_table_borders(table, color, 8)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    label_run = p.add_run(f"{label}  ")
    set_run_font(label_run, size=10, color=color, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_definition_table(doc, rows, widths=(2700, 6660), header=None):
    count = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=count, cols=2)
    set_table_geometry(table, list(widths))
    set_table_borders(table)
    row_offset = 0
    if header:
        set_repeat_table_header(table.rows[0])
        for idx, text in enumerate(header):
            cell = table.cell(0, idx)
            shade_cell(cell, BLUE_FILL)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_run_font(run, size=10, color=BLUE_DARK, bold=True)
        row_offset = 1
    for row_idx, (label, value) in enumerate(rows, start=row_offset):
        label_cell = table.cell(row_idx, 0)
        value_cell = table.cell(row_idx, 1)
        shade_cell(label_cell, CREAM)
        p_label = label_cell.paragraphs[0]
        p_label.paragraph_format.space_after = Pt(0)
        run = p_label.add_run(label)
        set_run_font(run, size=10.5, color=ACCENT, bold=True)
        p_value = value_cell.paragraphs[0]
        p_value.paragraph_format.space_after = Pt(0)
        run = p_value.add_run(value)
        set_run_font(run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        shade_cell(cell, BLUE_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=9.5, color=BLUE_DARK, bold=True)
    for row_index, row in enumerate(rows, start=1):
        for col, text in enumerate(row):
            cell = table.cell(row_index, col)
            if row_index % 2 == 0:
                shade_cell(cell, CREAM)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(text))
            set_run_font(run, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_compact_list(doc, items, columns=2):
    rows = (len(items) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    widths = [PAGE_WIDTH_DXA // columns] * columns
    set_table_geometry(table, widths)
    set_table_borders(table, BORDER, 4)
    for idx in range(rows * columns):
        cell = table.cell(idx // columns, idx % columns)
        if idx < len(items):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            number = p.add_run(f"{idx + 1}. ")
            set_run_font(number, size=10.5, color=ACCENT, bold=True)
            value = p.add_run(items[idx])
            set_run_font(value, size=10.5, color=INK)
        else:
            shade_cell(cell, CREAM)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_text_columns(doc, left_label, left_items, right_label, right_items):
    table = doc.add_table(rows=2, cols=2)
    set_table_geometry(table, [4680, 4680])
    set_table_borders(table)
    for col, label in enumerate((left_label, right_label)):
        shade_cell(table.cell(0, col), BLUE_FILL)
        p = table.cell(0, col).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, size=10, color=BLUE_DARK, bold=True)
    for col, items in enumerate((left_items, right_items)):
        cell = table.cell(1, col)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        for idx, item in enumerate(items):
            run = p.add_run(("" if idx == 0 else "\n") + item)
            set_run_font(run, size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def make_pair_image(paths, output_name):
    canvas_w, canvas_h = 1800, 760
    gap = 28
    panel_w = (canvas_w - gap) // 2
    canvas = Image.new("RGB", (canvas_w, canvas_h), "white")
    for index, path in enumerate(paths):
        image = Image.open(path).convert("RGB")
        fitted = ImageOps.fit(
            image,
            (panel_w, canvas_h),
            method=Image.Resampling.LANCZOS,
            centering=(0.5, 0.5),
        )
        canvas.paste(fitted, (index * (panel_w + gap), 0))
    out = ASSETS / output_name
    canvas.save(out, quality=94, optimize=True)
    return out


def add_photo_pair(doc, paths, name, width=6.5, caption=True):
    image_path = make_pair_image(paths, name)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    if caption:
        caption_paragraph = doc.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.space_after = Pt(9)
        run = caption_paragraph.add_run("Bolsa sellada / presentación del mealprep")
        set_run_font(run, size=9, color=MUTED, italic=True)


def add_full_image(doc, path, width=6.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(width))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        run = cp.add_run(caption)
        set_run_font(run, size=9, color=MUTED, italic=True)


def add_nutrition_table(doc, meal):
    add_labeled_paragraph(doc, "Porción", meal["portion"])
    add_matrix_table(
        doc,
        ["Indicador", "Por 100 g", "Por porción"],
        meal["nutrition"],
        [4560, 2400, 2400],
    )


def add_validation_box(doc, title="Decisión de Cecilia"):
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=3, cols=1)
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    set_table_borders(table)
    choices = [
        "☐ Sí, corresponde y se puede conservar como ficha válida.",
        "☐ No corresponde; no usar esta información.",
        "☐ Corregir antes de usar. Observación: __________________________________________",
    ]
    for idx, text in enumerate(choices):
        cell = table.cell(idx, 0)
        if idx == 0:
            shade_cell(cell, GREEN_LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=10.5, color=INK, bold=idx == 0)


def page_break(doc):
    global PENDING_PAGE_BREAK
    PENDING_PAGE_BREAK = True


def build_document():
    doc = Document()
    configure_document(doc)

    # Page 1: customer_pack opening.
    add_kicker(doc, "Documento para Cecilia")
    add_title(doc, "Validación de recuperación de mealpreps")
    add_subtitle(
        doc,
        "Contenido encontrado en Supabase, R2 y una grabación del 25 de julio de 2026.",
    )
    add_definition_table(
        doc,
        [
            ("Preparado para", "María Cecilia Salas · Fullness Lab"),
            ("Fecha de corte", "28 de julio de 2026"),
            ("Objetivo", "Confirmar qué contenido corresponde antes de recuperarlo"),
            ("Estado", "Sólo lectura · sin restauraciones en producción"),
        ],
    )
    add_status_callout(
        doc,
        "IMPORTANTE",
        "No se ha escrito ni restaurado ningún dato en Supabase. Este documento sirve para validar primero y recuperar después.",
        RED_LIGHT,
        RED,
    )
    add_full_image(
        doc,
        R2 / "4de64733-36b1-4e88-9428-37b8c7d4c9f1.png",
        width=5.75,
        caption="Imagen recuperada de R2: caja semanal con seis mealpreps.",
    )

    page_break(doc)

    # Page 2: definitive language and recovery map.
    add_kicker(doc, "Conceptos definitivos")
    add_title(doc, "La estructura que usará Fullness")
    add_definition_table(
        doc,
        [
            (
                "Mealprep familiar",
                "Formato familiar vendido directamente. Antes se mostraba como “plato familiar”.",
            ),
            (
                "Planes",
                "Caja semanal o mensual compuesta por varios mealpreps.",
            ),
            (
                "Mealpreps",
                "Preparaciones individuales en bolsa sellada, listas para calentar a baño María y usar dentro de los planes.",
            ),
        ],
    )
    add_heading(doc, "Mapa de recuperación", level=1)
    add_matrix_table(
        doc,
        ["Contenido", "Estado", "Evidencia", "Acción propuesta"],
        [
            (
                "Plan semanal antinflamatorio",
                "Vigente",
                "Supabase + video + R2",
                "Conservar; validar precio y seis nombres",
            ),
            (
                "Apple Golden Chicken",
                "Completo",
                "Supabase + video + R2",
                "Conservar; revisar variante histórica",
            ),
            (
                "Merluza y Crispy Quinoa",
                "Completo",
                "Supabase + R2",
                "Crear vínculo al catálogo sólo tras aprobación",
            ),
            (
                "Crepes de espinaca",
                "Parcial",
                "Nombre + dos imágenes R2",
                "Completar con Cecilia; no inventar datos",
            ),
            (
                "Otros tres mealpreps",
                "Nombre",
                "Descripción del plan + caja R2",
                "Crear fichas nuevas, no restaurarlas",
            ),
            (
                "Contenido demostrativo antiguo",
                "Excluido",
                "Código y capturas previas",
                "No restaurar",
            ),
        ],
        [2600, 1300, 2450, 3010],
    )
    add_status_callout(
        doc,
        "CRITERIO",
        "Una imagen o un nombre demuestra que el mealprep existía, pero no autoriza a completar ingredientes, alérgenos o nutrición por inferencia.",
        AMBER_LIGHT,
        AMBER,
    )

    page_break(doc)

    # Page 3: plan.
    add_kicker(doc, "Plan recuperado")
    add_title(doc, "Plan semanal antinflamatorio")
    add_definition_table(
        doc,
        [
            ("Estado actual", "Activo en Supabase"),
            ("Frecuencia", "Semanal"),
            ("Precio vigente", "$58.200 CLP"),
            ("Formato", "6 mealpreps / 1 semana"),
            (
                "Preparación",
                "Calentar la bolsa sellada a baño María entre 8 y 12 minutos, según el producto.",
            ),
        ],
    )
    add_heading(doc, "Los seis mealpreps de la caja", level=2)
    add_compact_list(doc, PLAN_MEALPREPS, columns=2)
    add_heading(doc, "Qué demuestra la grabación", level=2)
    add_labeled_paragraph(
        doc,
        "25 de julio",
        (
            "Cecilia muestra la caja con seis mealpreps, una descripción general y un "
            "precio de $57.800. Explica que completó la primera ficha y que seguiría "
            "con las restantes hasta llegar a seis."
        ),
    )
    add_labeled_paragraph(
        doc,
        "28 de julio",
        (
            "Supabase conserva el plan con precio de $58.200, la descripción que "
            "enumera los seis mealpreps y dos fichas completas. El precio de $57.800 "
            "se considera una versión anterior y no debe reemplazar al vigente."
        ),
    )
    add_status_callout(
        doc,
        "CONCLUSIÓN",
        "El plan no está perdido. Hay que conservarlo y volver a enlazar, con aprobación, las fichas recuperables.",
        GREEN_LIGHT,
        GREEN,
    )

    page_break(doc)

    # Page 4: Apple main fiche.
    add_kicker(doc, "Candidato 1 · alta certeza")
    add_title(doc, APPLE["name"], size=23)
    add_status_callout(
        doc,
        APPLE["status"],
        "La ficha vive dentro del plan semanal y también aparece duplicada en el catálogo técnico. Sus imágenes siguen en R2.",
        GREEN_LIGHT,
        GREEN,
    )
    add_photo_pair(doc, APPLE["photos"], "apple-pair.jpg")
    add_labeled_paragraph(doc, "Descripción vigente", APPLE["description"])
    add_labeled_paragraph(doc, "Descripción nutricional", APPLE["nutrition_description"])
    add_text_columns(
        doc,
        "Ingredientes recuperados",
        APPLE["ingredients"],
        "Beneficios y alérgenos",
        [f"Beneficio: {value}" for value in APPLE["benefits"]]
        + [f"Alérgeno: {value}" for value in APPLE["allergens"]],
    )

    # Page 5: Apple nutrition and historical comparison.
    add_kicker(doc, "Candidato 1 · detalle")
    add_title(doc, "Apple Golden Chicken · nutrición", size=23)
    add_nutrition_table(doc, APPLE)
    add_labeled_paragraph(
        doc,
        "Características",
        " · ".join(APPLE["highlights"]),
    )
    add_heading(doc, "Variante histórica visible en el video", level=2)
    add_labeled_paragraph(
        doc,
        "Descripción parcial",
        (
            "“Pollo marinado en especias aromáticas, manzana verde y leche de "
            "almendras, cocinado lentamente para lograr una textura tierna y jugosa…”"
        ),
    )
    add_labeled_paragraph(
        doc,
        "Advertencia de alérgenos",
        (
            "En la grabación también se ve una advertencia más amplia de elaboración "
            "en cocina compartida con pescado, lácteos, sésamo, huevo, frutos secos y "
            "legumbres. Debe confirmarse antes de reemplazar el dato vigente “Almendras”."
        ),
    )
    add_status_callout(
        doc,
        "NO APLICAR AUTOMÁTICAMENTE",
        "La variante del video es anterior y está parcialmente visible. Se conserva como evidencia para que Cecilia elija la redacción correcta.",
        AMBER_LIGHT,
        AMBER,
    )
    add_validation_box(doc)

    page_break(doc)

    # Page 6: Merluza main fiche.
    add_kicker(doc, "Candidato 2 · alta certeza")
    add_title(doc, MERLUZA["name"], size=23)
    add_status_callout(
        doc,
        MERLUZA["status"],
        "La ficha completa permanece dentro del plan semanal y las dos imágenes están disponibles en R2.",
        GREEN_LIGHT,
        GREEN,
    )
    add_photo_pair(doc, MERLUZA["photos"], "merluza-pair.jpg", width=5.4, caption=False)
    add_labeled_paragraph(doc, "Descripción vigente", MERLUZA["description"])
    add_heading(doc, "Ingredientes recuperados", level=2)
    add_compact_list(doc, MERLUZA["ingredients"], columns=3)
    add_labeled_paragraph(doc, "Beneficios recuperados", " · ".join(MERLUZA["benefits"]))

    page_break(doc)

    # Page 7: Merluza nutrition and validation.
    add_kicker(doc, "Candidato 2 · detalle")
    add_title(doc, "Merluza y Crispy Quinoa · nutrición", size=23)
    add_labeled_paragraph(doc, "Descripción nutricional", MERLUZA["nutrition_description"])
    add_nutrition_table(doc, MERLUZA)
    add_labeled_paragraph(
        doc,
        "Características",
        " · ".join(MERLUZA["highlights"]),
    )
    add_heading(doc, "Alérgenos recuperados", level=2)
    for allergen in MERLUZA["allergens"]:
        add_labeled_paragraph(doc, "Revisar", allergen, after=3)
    add_validation_box(doc)

    page_break(doc)

    # Page 8: Crepes and name-only candidates.
    add_kicker(doc, "Candidato 3 · certeza parcial")
    add_title(doc, CREPES["name"], size=17)
    add_status_callout(
        doc,
        CREPES["status"],
        "Se recuperaron el nombre y dos imágenes. No aparecieron una ficha nutricional, ingredientes ni alérgenos confiables.",
        AMBER_LIGHT,
        AMBER,
    )
    add_photo_pair(doc, CREPES["photos"], "crepes-pair.jpg", width=5.0, caption=False)
    add_definition_table(
        doc,
        [
            ("Confirmado", "Nombre en la descripción vigente del plan"),
            ("Confirmado", "Bolsa sellada y fotografía del mealprep en R2"),
            ("No recuperado", "Descripción, ingredientes, alérgenos, beneficios y nutrición"),
            ("Acción segura", "Cecilia completa la ficha; no se infiere contenido desde la foto"),
        ],
    )
    add_heading(doc, "Mealpreps con nombre, pero sin ficha recuperada", level=2)
    add_labeled_paragraph(doc, "Pendientes", " · ".join(PLAN_MEALPREPS[3:]))
    add_validation_box(doc)

    page_break(doc)

    # Page 9: safe Supabase approach and final approval.
    add_kicker(doc, "Siguiente paso")
    add_title(doc, "Recuperar sin tocar lo avanzado")
    add_heading(doc, "Estado de Supabase", level=1)
    add_labeled_paragraph(
        doc,
        "Datos actuales",
        (
            "El proyecto conserva el plan semanal, dos mealpreps completos, dos "
            "mealpreps familiares y un plan mensual. La tabla remota de borradores "
            "no existe, por lo que esos borradores no quedaron sincronizados allí."
        ),
    )
    add_labeled_paragraph(
        doc,
        "Versiones anteriores",
        (
            "La sesión de propietario confirmó que Fullness está en el plan Free. "
            "Supabase indica “No backups” y que este plan no incluye respaldos del "
            "proyecto ni recuperación a un punto anterior. No existe una versión "
            "histórica que podamos clonar o restaurar."
        ),
    )
    add_labeled_paragraph(
        doc,
        "Registro de actividad",
        (
            "También se revisaron los eventos disponibles para menu_items, "
            "meal_library_items y backoffice_drafts. El registro conserva rutas, "
            "métodos y estados, pero no el contenido enviado por Cecilia; por eso no "
            "permite reconstruir las fichas perdidas."
        ),
    )
    add_heading(doc, "Flujo seguro propuesto", level=2)
    add_matrix_table(
        doc,
        ["Paso", "Acción", "Protección"],
        [
            ("1", "Exportar una copia de las tablas actuales", "Punto de retorno verificable"),
            ("2", "Cecilia valida este dossier campo por campo", "No se infiere contenido"),
            ("3", "Preparar sólo los registros aprobados", "Lo avanzado permanece intacto"),
            ("4", "Aplicar una recuperación aditiva y auditada", "Sin reemplazos masivos"),
            ("5", "Activar borradores remotos y respaldos futuros", "Prevención permanente"),
        ],
        [720, 5040, 3600],
    )
    add_status_callout(
        doc,
        "REGLA",
        "No se escribirá ninguna recuperación hasta que Cecilia apruebe este dossier. La recuperación se suma a los datos actuales; no los reemplaza.",
        RED_LIGHT,
        RED,
    )

    page_break(doc)
    add_kicker(doc, "Cierre de validación")
    add_title(doc, "Aprobación general", size=24)
    add_validation_box(doc, title="Decisión sobre este dossier")
    add_labeled_paragraph(
        doc,
        "Observaciones",
        "________________________________________________________________________________\n________________________________________________________________________________",
    )
    add_heading(doc, "Fuentes", level=3)
    for source in [
        "Supabase actual · lectura del 28-07-2026",
        "Cloudflare R2 · 33 objetos revisados, 10 imágenes únicas",
        "Video de Cecilia · 25-07-2026, duración 1:31",
        "Historial del proyecto y migraciones locales",
        "Supabase Database Backups · https://supabase.com/features/database-backups",
        "Supabase Restore to a New Project · https://supabase.com/docs/guides/platform/clone-project",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(source)
        set_run_font(run, size=9, color=MUTED)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
